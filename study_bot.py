import streamlit as st
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
import docx
import tempfile
import requests
import json
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from uuid import uuid4
from langchain_core.documents import Document
from dotenv import load_dotenv

load_dotenv()

embeddings = OpenAIEmbeddings(model="text-embedding-3-large")

vector_store = Chroma(
    collection_name="example_collection",
    embedding_function=embeddings,
    persist_directory="./local_vector_store",  # Where to save data locally, remove if not necessary
)

def extract_text_with_ocr(file):
    """Use OCR to extract text from a scanned PDF."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(file.read())
        temp_path = temp_file.name

    images = convert_from_path(temp_path)
    text = ""
    for img in images:
        text += pytesseract.image_to_string(img)
    return text

def extract_text_directly(file):
    """Extract text from a PDF file directly."""
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        text += page_text if page_text else ""
    return text

def extract_text_from_docx(file):
    """Extract text from a Word document."""
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def filter_common_questions(text_list):
    """Filter out duplicate questions from a list."""
    question_set = set()
    unique_questions = []
    for question in text_list:
        if question.strip() not in question_set:
            question_set.add(question.strip())
            unique_questions.append(question.strip())
    return unique_questions

def get_chatbot_response_questions(user_message, question):
    """Get a response from the chatbot API with context."""
    url = "https://api.worqhat.com/api/ai/content/v4"
    
    # Add context (previous extracted questions) to the prompt
    question_text = str(question)
    
    payload = json.dumps({
        "question": user_message + question_text,
        "model": "aicon-v4-large-160824",
        "randomness": 0.1,
        "stream_data": False,
        "training_data": f"You are an expert teacher",
        "response_type": "text"
    })
    
    headers = {
        'Content-Type': 'application/json',
        "Authorization": "Bearer sk-128fd96fdf0d4039bdaef731e091de03"  # Replace with your token
    }
    
    response = requests.post(url, headers=headers, data=payload)

    
    if response.status_code == 200:
        
        return response.json().get("content", "No response received.")
    else:
        return f"Error: {response.status_code}, {response.text}"

def get_chatbot_response_answers(user_message, question):
    """Get a response from the chatbot API with context."""
    url = "https://api.worqhat.com/api/ai/content/v4"
    
    # Add context (previous extracted questions) to the prompt
    question_text = str(question)
    
    payload = json.dumps({
        "question": user_message + question_text,
        "model": "aicon-v4-large-160824",
        "randomness": 0.1,
        "stream_data": False,
        "training_data": f"You are an expert teacher",
        "response_type": "text"
    })
    
    headers = {
        'Content-Type': 'application/json',
        "Authorization": "Bearer sk-128fd96fdf0d4039bdaef731e091de03"  # Replace with your token
    }
    
    response = requests.post(url, headers=headers, data=payload)

    
    if response.status_code == 200:
        print("QUESTION", "\n", question_text)
        answer = response.json().get("content")
        print("ANSWER", "\n", answer)

        document = Document(
            page_content = question_text,
            metadata={"answer": answer},
            id= uuid4(),
        )
        vector_store.add_documents([document])
        print("Document added:", document.metadata)
        return response.json().get("content", "No response received.")
    else:
        return f"Error: {response.status_code}, {response.text}"



def get_chatbot_response_with_context(message, questions):
    """Send multiple API requests with smaller chunks of context and display each response."""
    responses = []
    


    # Split the context into chunks of size 'max_chunk_size'
    for index , question in enumerate(questions):
        print(question)
        response = get_chatbot_response_answers(message, question)
        responses.append(response)
        st.write(f"ANSWER {index} \n {response}")  # Display each response immediately

        st.divider()
    
    # Combine all the responses from each chunk
    full_response = "\n".join(responses)
    return full_response

# def storeInChroma()


def main():
    st.title("TARS bot")
  

    # File Upload Section
    uploaded_files = st.file_uploader(
        "Upload question papers (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=True
    )

    extracted_questions = []

    if uploaded_files:
        # Define available options
        options = ["Direct Extraction (for selectable text)", "OCR (for scanned documents)"]

        # Simulate a condition to disable "OCR"
        disable_ocr = True

        if disable_ocr:
            # Indicate the option is disabled
            options = ["Direct Extraction (for selectable text)"]  # Remove OCR or mark it disabled
            st.warning("OCR extraction is coming soon")

        # Display radio button
        extraction_mode = st.radio(
            "Choose Extraction Method:",
            options
        )

        
        

        for file in uploaded_files:
            file_name = file.name
          

            if file_name.endswith(".pdf"):
                if extraction_mode == "Direct Extraction (for selectable text)":
                    text = extract_text_directly(file)
                else:
                    text = extract_text_with_ocr(file)
            elif file_name.endswith(".docx"):
                text = extract_text_from_docx(file)
            else:
                st.error("Unsupported file format.")
                continue

            questions = text.split("\n")  # Split text into lines
            extracted_questions.extend(questions)

    prompt = """
    This is the text extracted from a file. Your task is to extract and format only the questions from the text. Follow these rules strictly:

Questions are indicated by "Q<question number>)" (e.g., Q1) What is your name?).
Subquestions are marked by letters or Roman numerals (e.g., a), b), i), ii), etc.).
If a subquestion contains multiple parts (e.g., i) Compiler, ii) Interpreter), combine them under the same subquestion heading as a single item. Do not split them into separate questions.
Write each subquestion on a new line.
Do not combine multiple main questions into one.
Do not include any additional text, context, or comments outside of the questions. Do not put new line at the end or after question number. I want one subquestion per line.
Example Input:
a) Write short note on [7]
i) Compiler
ii) Interpreter

Example Output:
a) Write a short note on: Compiler and Interpreter

Here is the text:
"""
    st.divider()
    if extracted_questions:
        st.header("Get the AI to frame the questions properly.")
        if st.button("Frame Questions"):
            st.write("Framing questions from the uploaded content...")
          
            framed_questions = get_chatbot_response_questions(prompt, extracted_questions).split("\n")
            st.session_state["framed_questions"] = framed_questions  # Store framed questions in session state


    # Display framed questions for confirmation
    if "framed_questions" in st.session_state:
        st.subheader("Framed Questions (Confirm or Edit)")
        confirmed_questions = []
        for idx, question in enumerate(st.session_state["framed_questions"]):
            user_input = st.text_input(f"Question {idx + 1}:", value=question)
            confirmed_questions.append(user_input)
        st.divider()
        st.session_state["confirmed_questions"] = confirmed_questions

 

    # Chatbot Section
    st.subheader("Answers (They will appear sequentially as the bot responds.)")
    user_question = st.text_input("Enter context about the subject")

    if st.button("Get all answers"):
        if user_question:
            st.write(confirmed_questions)

            # Collect questions to remove in a separate list
            questions_to_remove = []

            for index, question in enumerate(confirmed_questions):
                results = vector_store.similarity_search_with_score(
                    question,
                    k=1,
                )

                st.write("Checking question: ", question)
                found = False  # Flag to indicate if the question is found

                for res, score in results:
                    similarity_score = score
                    print(f"* [SIM={score:3f}] {res.page_content}")
                    if similarity_score < 0.001:
                        st.write(res.metadata["answer"])
                        questions_to_remove.append(question)  # Add to removal list
                       
                        found = True
                        break

                if not found:
                    st.write("NOT FOUND: ", question)

            # Remove the collected questions after iteration
            for question in questions_to_remove:
                confirmed_questions.remove(question)

            st.write(f"Remaining questions: {confirmed_questions}")
            full_response = get_chatbot_response_with_context(user_question, confirmed_questions)

if __name__ == "__main__":
    main()
